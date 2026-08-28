"""DOCX renderer.

Layout is a single column with a hanging indent, so the date sits in a left
gutter and the body block-aligns beside it.  A borderless table would also
work but splits badly across pages; a hanging indent lets Word's own
``keepLines`` / ``keepNext`` keep an entry and its heading together, which is
what a long publication list needs.

python-docx has no hyperlink API, so ``add_hyperlink`` builds the w:hyperlink
element and the relationship by hand.  Relationship ids are allocated
sequentially, which keeps the output deterministic.
"""

from __future__ import annotations

import zipfile
from datetime import datetime
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm, Pt, RGBColor

from .model import Document, Entry, Run, Section
from .theme import (DATE_COL, FOOTER, FONT, PAGE, PHOTO,
                    RULE_COLOR, SIZE, SPACE)

REPO = Path(__file__).resolve().parents[2]

HYPERLINK_RT = ("http://schemas.openxmlformats.org/officeDocument/2006/"
                "relationships/hyperlink")
#: fixed timestamp so repeated builds are byte-identical
EPOCH = datetime(2000, 1, 1)


# --------------------------------------------------------------------------
# low-level helpers
# --------------------------------------------------------------------------

def _set_run_fonts(run, *, size=None, bold=False, italic=False,
                   color=None, cjk=True):
    run.font.name = FONT["latin"]
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT["latin"])
    rfonts.set(qn("w:hAnsi"), FONT["latin"])
    # East-Asian glyphs come from the CJK face; without this Word/LibreOffice
    # pick an arbitrary fallback and the Japanese CV loses its typeface.
    if cjk:
        rfonts.set(qn("w:eastAsia"), FONT["cjk"])


def add_hyperlink(paragraph, text: str, url: str, *, size=None, bold=False,
                  italic=False):
    part = paragraph.part
    r_id = part.relate_to(url, HYPERLINK_RT, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    link.append(new_run)
    paragraph._p.append(link)
    run = docx.text.run.Run(new_run, paragraph)
    run.text = text
    _set_run_fonts(run, size=size, bold=bold, italic=italic)
    # keep links black and undecorated: this is an academic CV, not a web page
    return run


def emit_runs(paragraph, runs, *, base_size: float):
    for r in runs:
        size = base_size * 0.92 if r.small else base_size
        if r.href:
            add_hyperlink(paragraph, r.text, r.href, size=size,
                          bold=r.bold, italic=r.italic)
        else:
            run = paragraph.add_run(r.text)
            _set_run_fonts(run, size=size, bold=r.bold, italic=r.italic)


def _page_field(paragraph):
    """Insert a { PAGE } field; python-docx has no API for this."""
    run = paragraph.add_run()
    _set_run_fonts(run, size=FOOTER["size"])
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, end):
        run._element.append(el)


def add_footer(section, lang: str, footer_style=None):
    """`S Namba | Curriculum Vitae` on the left, the page number on the right."""
    # The built-in Footer STYLE carries centre and right tab stops sized for
    # US Letter. Clearing them on the paragraph does not remove style-level
    # stops, so a single tab lands on the style's centre stop instead of our
    # right one -- clear the style itself first.
    try:
        style_tabs = footer_style.paragraph_format.tab_stops
        while len(style_tabs):
            style_tabs.clear_all()
            break
    except Exception:                                  # pragma: no cover
        pass
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.tab_stops.clear_all()
    usable = PAGE["width"] - PAGE["margin_left"] - PAGE["margin_right"]
    pf.tab_stops.add_tab_stop(Mm(usable), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(FOOTER["left"][lang])
    _set_run_fonts(run, size=FOOTER["size"])
    tab = p.add_run("\t")
    _set_run_fonts(tab, size=FOOTER["size"])
    _page_field(p)


def _borderless(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def _bottom_border(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")          # eighths of a point
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), RULE_COLOR)
    borders.append(bottom)
    ppr.append(borders)


# --------------------------------------------------------------------------
# renderer
# --------------------------------------------------------------------------

class DocxRenderer:
    def __init__(self, doc: Document):
        self.doc = doc
        self.lang = doc.language
        self.date_col = DATE_COL[self.lang]

    # -- paragraph builders ------------------------------------------------
    def _p(self, *, space_before=0.0, space_after=0.0, keep_next=False,
           keep_together=True):
        p = self.docx.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = SPACE["line"]
        pf.keep_together = keep_together
        pf.keep_with_next = keep_next
        pf.widow_control = True
        return p

    def _heading(self, text: str):
        p = self._p(space_before=SPACE["before_section"],
                    space_after=SPACE["after_section"], keep_next=True)
        run = p.add_run(text)
        _set_run_fonts(run, size=SIZE["section"], bold=True)
        _bottom_border(p)

    def _entry(self, e: Entry, *, indent_mm: float, base_size: float,
               lead: str | None = None):
        p = self._p(space_after=SPACE["after_entry"])
        pf = p.paragraph_format
        pf.left_indent = Mm(indent_mm)
        pf.first_line_indent = Mm(-indent_mm)
        pf.tab_stops.add_tab_stop(Mm(indent_mm), WD_TAB_ALIGNMENT.LEFT)
        if e.date:
            emit_runs(p, [Run(r.text, small=True) for r in e.date],
                      base_size=SIZE["date"] / 0.92)
            p.add_run("\t")
        elif lead is not None:
            run = p.add_run(lead)
            _set_run_fonts(run, size=base_size)
            p.add_run("\t")
        emit_runs(p, list(e.body), base_size=base_size)
        if e.detail:
            d = self._p(space_after=SPACE["after_entry"])
            dpf = d.paragraph_format
            dpf.left_indent = Mm(indent_mm)
            # Same size as the body: the advisor's name should not look like
            # a footnote.
            emit_runs(d, list(e.detail), base_size=SIZE["body"])

    # -- sections ----------------------------------------------------------
    ROLE_SIZE = {"furigana": SIZE["contact"], "name": SIZE["name"],
                 "subtitle": SIZE["subtitle"], "contact": SIZE["contact"]}

    def _header_entry_sizes(self, s: Section):
        roles = self.doc.meta.get("header_roles") or []
        for i, e in enumerate(s.entries):
            role = roles[i] if i < len(roles) else "contact"
            yield e, self.ROLE_SIZE[role], role

    def _header_section(self, s: Section):
        photo = self.doc.meta.get("photo")
        photo_path = (REPO / photo) if photo else None
        if photo_path and photo_path.exists():
            self._header_with_photo(s, photo_path)
        else:
            self._header_plain(s)
        self._p(space_after=SPACE["after_header"], keep_together=True)

    def _header_with_photo(self, s: Section, photo_path):
        """Name/affiliation block on the left, portrait on the right."""
        from PIL import Image
        with Image.open(photo_path) as im:
            w, h = im.size
        height = PHOTO["height_mm"]
        width = height * (w / h)
        usable = PAGE["width"] - PAGE["margin_left"] - PAGE["margin_right"]

        table = self.docx.add_table(rows=1, cols=2)
        _borderless(table)
        table.autofit = False
        left, right = table.rows[0].cells
        left.width = Mm(usable - width - PHOTO["gap_mm"])
        right.width = Mm(width + PHOTO["gap_mm"])

        first = True
        for e, size, role in self._header_entry_sizes(s):
            p = left.paragraphs[0] if first else left.add_paragraph()
            first = False
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(SPACE["after_name"] if role == "name" else 1.0)
            pf.line_spacing = SPACE["line"]
            emit_runs(p, list(e.body), base_size=size)

        cell_p = right.paragraphs[0]
        cell_p.paragraph_format.space_before = Pt(0)
        cell_p.paragraph_format.space_after = Pt(0)
        cell_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell_p.add_run().add_picture(str(photo_path), height=Mm(height))

    def _header_plain(self, s: Section):
        for e, size, role in self._header_entry_sizes(s):
            p = self._p(space_after=SPACE["after_name"] if role == "name"
                        else 1.0, keep_next=True, keep_together=True)
            emit_runs(p, list(e.body), base_size=size)

    def _paragraph_section(self, s: Section):
        self._heading(s.heading)
        for e in s.entries:
            p = self._p(space_after=SPACE["after_entry"])
            emit_runs(p, list(e.body), base_size=SIZE["body"])

    def _divider_section(self, s: Section):
        """Page break, then a title announcing the complete record."""
        p = self._p(space_before=0, space_after=SPACE["after_section"],
                    keep_next=True)
        p.paragraph_format.page_break_before = True
        run = p.add_run(s.heading)
        _set_run_fonts(run, size=SIZE["name"] * 0.75, bold=True)
        _bottom_border(p)

    def _plainlist_section(self, s: Section):
        self._heading(s.heading)
        for e in s.entries:
            p = self._p(space_after=1.0)
            emit_runs(p, list(e.body), base_size=SIZE["body"])

    def _entries_section(self, s: Section):
        self._heading(s.heading)
        for e in s.entries:
            self._entry(e, indent_mm=self.date_col, base_size=SIZE["body"])

    def _numbered_section(self, s: Section):
        self._heading(s.heading)
        if s.note:
            p = self._p(space_after=SPACE["after_entry"])
            emit_runs(p, list(s.note), base_size=SIZE["note"])
        for e in s.entries:
            self._entry(e, indent_mm=8.0, base_size=SIZE["body"])
        for sub in s.subsections:
            p = self._p(space_before=6, space_after=2, keep_next=True)
            run = p.add_run(sub.heading)
            _set_run_fonts(run, size=SIZE["body"], bold=True, italic=True)
            for e in sub.entries:
                self._entry(e, indent_mm=8.0, base_size=SIZE["body"])

    # -- entry point -------------------------------------------------------
    def render(self, out: Path) -> Path:
        self.docx = docx.Document()
        sec = self.docx.sections[0]
        sec.page_width = Mm(PAGE["width"])
        sec.page_height = Mm(PAGE["height"])
        sec.top_margin = Mm(PAGE["margin_top"])
        sec.bottom_margin = Mm(PAGE["margin_bottom"])
        sec.left_margin = Mm(PAGE["margin_left"])
        sec.right_margin = Mm(PAGE["margin_right"])

        style = self.docx.styles["Normal"]
        style.font.name = FONT["latin"]
        style.font.size = Pt(SIZE["body"])
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), FONT["cjk"])

        add_footer(sec, self.lang, self.docx.styles["Footer"])

        for s in self.doc.sections:
            if s.kind == "header":
                self._header_section(s)
            elif s.kind == "paragraph":
                self._paragraph_section(s)
            elif s.kind == "numbered":
                self._numbered_section(s)
            elif s.kind == "plainlist":
                self._plainlist_section(s)
            elif s.kind == "divider":
                self._divider_section(s)
            else:
                self._entries_section(s)

        cp = self.docx.core_properties
        cp.author = "Shinichi Namba"
        cp.title = f"CV ({self.doc.profile_name})"
        cp.created = EPOCH
        cp.modified = EPOCH
        cp.last_modified_by = ""
        cp.revision = 1

        out.parent.mkdir(parents=True, exist_ok=True)
        self.docx.save(out)
        normalize_zip(out)
        return out


def normalize_zip(path: Path) -> None:
    """Rewrite the archive so identical content produces identical bytes.

    python-docx stamps wall-clock times into the zip entries, which would make
    every rebuild differ and defeat the write-if-changed copy into assets/cv/.
    """
    with zipfile.ZipFile(path) as zf:
        items = sorted(zf.namelist())
        payload = {n: zf.read(n) for n in items}
    tmp = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
        for name in items:
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            info.create_system = 0
            zf.writestr(info, payload[name])
    tmp.replace(path)


def render_docx(doc: Document, out: Path) -> Path:
    return DocxRenderer(doc).render(out)
